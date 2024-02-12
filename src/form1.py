from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
def populate_docx_template(template_path, data):

    doc = Document(template_path)

    for table_itr, table in enumerate(doc.tables):
        if table_itr == 0:
            for row_itr, row in enumerate(table.rows):
                for i in range(len(row.cells)):
                    row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                if row_itr == 0:
                    row.cells[2].text = data['last_name'].upper()
                if row_itr == 1:
                    row.cells[2].text = data['first_name'].upper()
                if row_itr == 2:
                    row.cells[2].text = data['middle_name'].upper()

        if table_itr == 1:
             for row_itr, row in enumerate(table.rows):
                for i in range(len(row.cells)):
                    row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                if row_itr in [1, 2, 3]:
                    str_ = f'h_addr_line{row_itr}'
                    for i in range(4, 13):
                        row.cells[i].text = data[str_].upper()
                elif row_itr == 4:
                    row.cells[4].text = data['h_city'].upper()
                    row.cells[11].text = data['h_state'].upper()
                elif row_itr == 5:
                    row.cells[4].text = data['h_country'].upper()
                    row.cells[11].text = data['h_postal_code']
                elif row_itr == 8:
                    row.cells[4].text = data['p_addr_line1'].upper() 
                elif row_itr == 9:
                    row.cells[4].text = data['p_addr_line2'].upper() 
                elif row_itr == 10:
                    row.cells[4].text = data['p_addr_line3'].upper() 
                elif row_itr == 11:
                    row.cells[4].text = data['p_city'].upper()
                    row.cells[11].text = data['p_state'].upper()
                elif row_itr == 12:
                    row.cells[4].text = data['p_country'].upper()
                    row.cells[11].text = data['p_postal_code']
                elif row_itr == 17:
                    row.cells[3].text = data['mobile_c_code']
                    row.cells[8].text = data['mobile_number']
                elif row_itr == 21:
                    row.cells[1].text = data['email']
                else:
                    pass

        if table_itr == 2:
            for row_itr, row in enumerate(table.rows):
                for i in range(len(row.cells)):
                    row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                if row_itr == 0:
                    row.cells[2].text = data['gender'].upper()
                    row.cells[13].text = data['marital_status'].upper()
                    row.cells[16].text = data['marital_date']
                elif row_itr == 1:
                    row.cells[2].text = data['native_language'].upper()
                elif row_itr == 2:
                    row.cells[2].text = data['dob_dd']
                    row.cells[5].text = data['dob_mm']
                    row.cells[11].text = data['dob_yyyy']
                elif row_itr == 4:
                    row.cells[2].text = data['city_of_birth'].upper()
                    row.cells[12].text = data['country_of_birth'].upper()
                elif row_itr == 6:
                    row.cells[2].text = data['country_n_id1'].upper()
                    row.cells[7].text = data['type_n_id1'].upper()
                    row.cells[12].text = data['number_n_id1']
                elif row_itr == 8:
                    row.cells[2].text = data['country_n_id2'].upper()
                    row.cells[7].text = data['type_n_id2'].upper()
                    row.cells[12].text = data['number_n_id2']
                else:
                    pass
        
        if table_itr == 6:
            for row_itr, row in enumerate(table.rows):
                for i in range(len(row.cells)):
                    row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                if row_itr == 2:
                    row.cells[1].text = data['emergency_name'].upper()
                    row.cells[3].text = data['emergency_contact']

        for row in table.rows:
            for cell in row.cells:
                # Create a new paragraph and set its font size and name
                p = cell.paragraphs[0]
                for run in p.runs:
                    run.font.size = Pt(8)  # Set font size to 10 for table text

    return doc